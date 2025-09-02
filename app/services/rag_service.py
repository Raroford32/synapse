"""RAG Service for document processing and retrieval"""
import logging
from typing import List, Dict, Any, Optional, BinaryIO
from pathlib import Path
import hashlib
import mimetypes
from datetime import datetime

from sqlalchemy import select, and_, or_, func
from sqlalchemy.ext.asyncio import AsyncSession
import pypdf
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup

from app.models.document import Document, DocumentChunk, Collection, DocumentCollection
from app.services.embedding_service import EmbeddingService
from app.core.config import settings

logger = logging.getLogger(__name__)


class RAGService:
    """Service for document processing and retrieval"""
    
    SUPPORTED_FORMATS = {
        '.pdf': 'pdf',
        '.txt': 'text',
        '.md': 'markdown',
        '.docx': 'docx',
        '.html': 'html',
        '.py': 'code',
        '.js': 'code',
        '.ts': 'code',
        '.java': 'code',
        '.cpp': 'code',
        '.go': 'code',
    }
    
    def __init__(self, embedding_service: EmbeddingService):
        self.embedding_service = embedding_service
        self.chunk_size = settings.R2R_CHUNK_SIZE
        self.chunk_overlap = settings.R2R_CHUNK_OVERLAP
    
    async def ingest_document(
        self,
        session: AsyncSession,
        user_id: str,
        file_path: str,
        file_content: bytes,
        filename: str,
        metadata: Optional[Dict] = None
    ) -> Document:
        """Ingest a document and create chunks"""
        
        try:
            # Get file extension and type
            file_ext = Path(filename).suffix.lower()
            file_type = self.SUPPORTED_FORMATS.get(file_ext, 'unknown')
            
            if file_type == 'unknown':
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Extract text content
            content = await self._extract_content(file_content, file_type, filename)
            
            # Create document record
            document = Document(
                user_id=user_id,
                filename=filename,
                file_type=file_type,
                file_size=len(file_content),
                content=content,
                metadata=metadata or {},
                is_processed=False
            )
            
            session.add(document)
            await session.flush()  # Get document ID
            
            # Create chunks
            chunks = self._create_chunks(content)
            
            # Generate embeddings for chunks
            chunk_texts = [chunk['text'] for chunk in chunks]
            embeddings = await self.embedding_service.generate_embedding(chunk_texts)
            
            # Save chunks to database
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                chunk_record = DocumentChunk(
                    document_id=document.id,
                    chunk_index=i,
                    content=chunk['text'],
                    embedding=embedding,
                    metadata=chunk.get('metadata', {}),
                    token_count=len(chunk['text'].split()),
                    start_char=chunk['start'],
                    end_char=chunk['end']
                )
                session.add(chunk_record)
            
            # Mark document as processed
            document.is_processed = True
            
            await session.commit()
            await session.refresh(document)
            
            logger.info(f"Ingested document {filename} with {len(chunks)} chunks")
            return document
            
        except Exception as e:
            logger.error(f"Error ingesting document: {str(e)}")
            document.is_processed = False
            document.processing_error = str(e)
            await session.rollback()
            raise
    
    async def _extract_content(
        self,
        file_content: bytes,
        file_type: str,
        filename: str
    ) -> str:
        """Extract text content from various file formats"""
        
        try:
            if file_type == 'pdf':
                return self._extract_pdf(file_content)
            elif file_type == 'text' or file_type == 'code':
                return file_content.decode('utf-8', errors='ignore')
            elif file_type == 'markdown':
                text = file_content.decode('utf-8', errors='ignore')
                # Convert markdown to plain text
                html = markdown.markdown(text)
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text()
            elif file_type == 'docx':
                return self._extract_docx(file_content)
            elif file_type == 'html':
                soup = BeautifulSoup(file_content, 'html.parser')
                return soup.get_text()
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error extracting content from {filename}: {str(e)}")
            raise
    
    def _extract_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        
        import io
        text = []
        
        with io.BytesIO(file_content) as pdf_file:
            pdf_reader = pypdf.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        
        return '\n'.join(text)
    
    def _extract_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        
        import io
        doc = DocxDocument(io.BytesIO(file_content))
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        return '\n'.join(text)
    
    def _create_chunks(self, content: str) -> List[Dict]:
        """Create overlapping chunks from content"""
        
        chunks = []
        words = content.split()
        
        if not words:
            return chunks
        
        # Calculate chunk boundaries
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            start_idx = i
            end_idx = min(i + self.chunk_size, len(words))
            
            # Get chunk text
            chunk_words = words[start_idx:end_idx]
            chunk_text = ' '.join(chunk_words)
            
            # Calculate character positions
            start_char = len(' '.join(words[:start_idx])) + (1 if start_idx > 0 else 0)
            end_char = len(' '.join(words[:end_idx]))
            
            chunks.append({
                'text': chunk_text,
                'start': start_char,
                'end': end_char,
                'metadata': {
                    'chunk_index': len(chunks),
                    'word_count': len(chunk_words)
                }
            })
            
            # Stop if we've processed all words
            if end_idx >= len(words):
                break
        
        return chunks
    
    async def search_documents(
        self,
        session: AsyncSession,
        user_id: str,
        query: str,
        limit: int = 10,
        collection_id: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """Search documents using semantic search"""
        
        # Generate query embedding
        query_embedding = await self.embedding_service.generate_embedding(query)
        
        # Build base query
        chunk_query = select(DocumentChunk, Document).join(
            Document, DocumentChunk.document_id == Document.id
        ).where(Document.user_id == user_id)
        
        # Filter by collection if specified
        if collection_id:
            chunk_query = chunk_query.join(
                DocumentCollection,
                DocumentCollection.document_id == Document.id
            ).where(DocumentCollection.collection_id == collection_id)
        
        # Execute query
        result = await session.execute(chunk_query)
        chunks_with_docs = result.all()
        
        if not chunks_with_docs:
            return []
        
        # Calculate similarities
        results = []
        for chunk, doc in chunks_with_docs:
            if chunk.embedding:
                similarity = self.embedding_service.cosine_similarity(
                    query_embedding, chunk.embedding
                )
                
                results.append({
                    'document': doc,
                    'chunk': chunk,
                    'similarity': similarity,
                    'content': chunk.content,
                    'metadata': {
                        **doc.metadata,
                        **chunk.metadata,
                        'filename': doc.filename,
                        'chunk_index': chunk.chunk_index
                    }
                })
        
        # Sort by similarity
        results.sort(key=lambda x: x['similarity'], reverse=True)
        
        return results[:limit]
    
    async def get_document(
        self,
        session: AsyncSession,
        document_id: str,
        user_id: str
    ) -> Optional[Document]:
        """Get a specific document"""
        
        query = select(Document).where(
            and_(
                Document.id == document_id,
                Document.user_id == user_id
            )
        )
        
        result = await session.execute(query)
        return result.scalar_one_or_none()
    
    async def list_documents(
        self,
        session: AsyncSession,
        user_id: str,
        limit: int = 50,
        offset: int = 0
    ) -> List[Document]:
        """List user's documents"""
        
        query = select(Document).where(
            Document.user_id == user_id
        ).order_by(
            Document.created_at.desc()
        ).limit(limit).offset(offset)
        
        result = await session.execute(query)
        return result.scalars().all()
    
    async def delete_document(
        self,
        session: AsyncSession,
        document_id: str,
        user_id: str
    ) -> bool:
        """Delete a document and its chunks"""
        
        document = await self.get_document(session, document_id, user_id)
        
        if not document:
            return False
        
        await session.delete(document)
        await session.commit()
        
        logger.info(f"Deleted document {document_id} for user {user_id}")
        return True
    
    async def create_collection(
        self,
        session: AsyncSession,
        user_id: str,
        name: str,
        description: Optional[str] = None,
        metadata: Optional[Dict] = None
    ) -> Collection:
        """Create a document collection"""
        
        collection = Collection(
            user_id=user_id,
            name=name,
            description=description,
            metadata=metadata or {}
        )
        
        session.add(collection)
        await session.commit()
        await session.refresh(collection)
        
        return collection
    
    async def add_to_collection(
        self,
        session: AsyncSession,
        document_id: str,
        collection_id: str,
        user_id: str
    ) -> bool:
        """Add document to collection"""
        
        # Verify ownership
        doc = await self.get_document(session, document_id, user_id)
        if not doc:
            return False
        
        # Check if already in collection
        existing = await session.execute(
            select(DocumentCollection).where(
                and_(
                    DocumentCollection.document_id == document_id,
                    DocumentCollection.collection_id == collection_id
                )
            )
        )
        
        if existing.scalar_one_or_none():
            return True  # Already in collection
        
        # Add to collection
        doc_collection = DocumentCollection(
            document_id=document_id,
            collection_id=collection_id,
            added_by=user_id
        )
        
        session.add(doc_collection)
        
        # Update collection document count
        collection = await session.get(Collection, collection_id)
        if collection:
            collection.document_count += 1
        
        await session.commit()
        
        return True